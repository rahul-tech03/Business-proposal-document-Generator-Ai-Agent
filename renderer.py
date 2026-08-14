"""
Renders (title, sections) into a polished .docx using python-docx. The
only module that knows about document formatting -- the graph nodes hand
it plain data and don't know what the final file looks like.
"""
from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ACCENT_COLOR = RGBColor(0x1F, 0x4E, 0x79)
SUBTLE_GRAY = RGBColor(0x59, 0x59, 0x59)


def _slugify(text: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9]+", "_", text.strip().lower()).strip("_")
    return slug[:60] or "document"


def _maybe_render_list(doc, text: str) -> bool:
    lines = [ln.strip() for ln in text.split("\n") if ln.strip()]
    numbered = all(re.match(r"^\d+[\).]\s+", ln) for ln in lines) if lines else False
    bulleted = all(re.match(r"^[-*•]\s+", ln) for ln in lines) if lines else False

    if len(lines) >= 2 and (numbered or bulleted):
        for ln in lines:
            clean = re.sub(r"^(\d+[\).]|[-*•])\s+", "", ln)
            doc.add_paragraph(clean, style="List Number" if numbered else "List Bullet")
        return True
    return False


def render_document(title: str, doc_type: str, sections: dict[str, str],
                     assumptions: list[str], output_dir: str) -> str:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    heading = doc.add_paragraph()
    run = heading.add_run(title)
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = ACCENT_COLOR

    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run(
        f"{doc_type.replace('_', ' ').title()}  |  Prepared {date.today().strftime('%B %d, %Y')}"
    )
    sub_run.font.size = Pt(11)
    sub_run.italic = True
    sub_run.font.color.rgb = SUBTLE_GRAY

    rule = doc.add_paragraph()
    pPr = rule._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F4E79")
    pBdr.append(bottom)
    pPr.append(pBdr)

    if assumptions:
        note = doc.add_paragraph()
        note_run = note.add_run("Note — assumptions made by the agent:")
        note_run.bold = True
        note_run.font.size = Pt(10)
        note_run.font.color.rgb = SUBTLE_GRAY
        for a in assumptions:
            bullet = doc.add_paragraph(a, style="List Bullet")
            for r in bullet.runs:
                r.font.size = Pt(10)
                r.font.color.rgb = SUBTLE_GRAY
        doc.add_paragraph()

    for heading_text, body in sections.items():
        h = doc.add_heading(heading_text, level=1)
        for r in h.runs:
            r.font.color.rgb = ACCENT_COLOR
        if not _maybe_render_list(doc, body):
            p = doc.add_paragraph(body)
            p.paragraph_format.space_after = Pt(10)

    footer_p = doc.sections[0].footer.paragraphs[0]
    footer_p.text = "Generated automatically by the Autonomous Document Agent"
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in footer_p.runs:
        r.font.size = Pt(8)
        r.font.color.rgb = SUBTLE_GRAY

    Path(output_dir).mkdir(parents=True, exist_ok=True)
    file_path = str(Path(output_dir) / f"{_slugify(title)}.docx")
    doc.save(file_path)
    return file_path
